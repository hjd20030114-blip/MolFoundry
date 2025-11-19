# -*- coding: utf-8 -*-
"""
生成《PRRSV 抑制剂设计平台》项目白皮书（Word 版）
- 自动汇总最近一次 run 的数据与图表
- 详尽介绍项目功能、模型原理与创新点、工程设计、实验数据与可视化
- 输出: results/<run>/reports/PRRSV_Whitepaper.docx （并拷贝一份到 HJD/reports/ 便于查阅）

运行:
    pip install -r HJD/requirements.txt
    python HJD/scripts/generate_whitepaper.py

可选参数:
    --run <run_dir_name>   指定具体 run 目录名（如: run_20250918_008）
"""

import os
import sys
import glob
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple

import pandas as pd
import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as e:
    print("[ERROR] 需要安装 python-docx: pip install python-docx")
    raise

PROJECT_ROOT = Path(__file__).resolve().parents[1]
RESULTS_DIR = PROJECT_ROOT / "results"
REPORT_OUT_DIR = PROJECT_ROOT / "HJD" / "reports"
REPORT_OUT_DIR.mkdir(parents=True, exist_ok=True)

# 备用图表目录（仓库内示例图）
EXPERIMENT_REPORT_DIR = PROJECT_ROOT / "HJD" / "experiment_report"


def find_latest_run_dir(explicit: Optional[str] = None) -> Optional[Path]:
    if explicit:
        p = RESULTS_DIR / explicit
        return p if p.exists() and p.is_dir() else None
    candidates = [d for d in RESULTS_DIR.glob("run_*") if d.is_dir()]
    if not candidates:
        return None
    return max(candidates, key=lambda x: x.stat().st_mtime)


def first_existing(paths: List[Path]) -> Optional[Path]:
    for p in paths:
        if p and p.exists():
            return p
    return None


def load_csv_safe(p: Optional[Path]) -> pd.DataFrame:
    try:
        if p and p.exists():
            return pd.read_csv(p)
    except Exception as e:
        print(f"[WARN] 读取CSV失败: {p}: {e}")
    return pd.DataFrame()


def collect_run_data(run_dir: Path) -> Dict:
    lig_dir = run_dir / "ligands"
    dock_dir = run_dir / "docking"
    admet_dir = run_dir / "admet"
    viz2d_dir = run_dir / "visualization_2d"
    viz3d_dir = run_dir / "visualization_3d"
    reports_dir = run_dir / "reports"
    reports_dir.mkdir(exist_ok=True)

    # 选择配体CSV（优先顺序）
    lig_csv = first_existing([
        lig_dir / "generated_ligands.csv",
        lig_dir / "dl_phase3_optimized_molecules.csv",
        lig_dir / "dl_phase2_generated_molecules.csv",
    ])

    # 对接与 ADMET
    dock_csv = first_existing([
        dock_dir / "docking_results.csv",
    ])
    admet_csv = first_existing([
        admet_dir / "admet_results.csv",
    ])

    lig_df = load_csv_safe(lig_csv)
    dock_df = load_csv_safe(dock_csv)
    admet_df = load_csv_safe(admet_csv)

    # 合并：优先以对接结果为主，若无则以配体为主
    if not dock_df.empty:
        base_df = dock_df
    else:
        base_df = lig_df

    merged_df = base_df.copy()
    if not merged_df.empty and not admet_df.empty:
        if "smiles" in merged_df.columns and "smiles" in admet_df.columns:
            # 避免重复列
            cols_to_use = [c for c in admet_df.columns if c != "compound_id"]
            merged_df = merged_df.merge(admet_df[cols_to_use], on="smiles", how="left")

    # 图片收集
    grid_png = None
    try:
        grid_candidates = list((viz2d_dir / "grids").glob("Top_*_Molecules_grid.png"))
        grid_png = max(grid_candidates, key=lambda x: x.stat().st_mtime) if grid_candidates else None
    except Exception:
        grid_png = None

    dash_png = viz3d_dir / "interactive_dashboard.png"
    if not dash_png.exists():
        dash_png = None

    # 备用图（仓库内示例）
    fallback_imgs = {
        "affinity": EXPERIMENT_REPORT_DIR / "binding_affinity_distribution.png",
        "top10": EXPERIMENT_REPORT_DIR / "top_10_molecules.png",
        "lipinski": EXPERIMENT_REPORT_DIR / "lipinski_compliance.png",
        "admet_props": EXPERIMENT_REPORT_DIR / "admet_properties.png",
        "grid50": EXPERIMENT_REPORT_DIR / "visualization_2d" / "grids" / "Top_50_Molecules_grid.png",
    }

    return {
        "lig_df": lig_df,
        "dock_df": dock_df,
        "admet_df": admet_df,
        "merged_df": merged_df,
        "grid_png": grid_png,
        "dash_png": dash_png,
        "fallback_imgs": fallback_imgs,
        "reports_dir": reports_dir,
    }


def safe_add_picture(document: Document, img_path: Optional[Path], width_in: float = 6.0, caption: str = ""):
    if img_path and img_path.exists():
        try:
            document.add_picture(str(img_path), width=Inches(width_in))
            if caption:
                p = document.add_paragraph(caption)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            document.add_paragraph(f"[图片插入失败] {img_path.name}: {e}")
    else:
        if caption:
            document.add_paragraph(f"[缺少图片] {caption} ({img_path})")


def summarize_numbers(df: pd.DataFrame) -> Dict:
    if df.empty:
        return {}
    out: Dict[str, object] = {}
    # 结合亲和力
    if "binding_affinity" in df.columns:
        vals = df["binding_affinity"].dropna()
        if len(vals) > 0:
            out["best_affinity"] = float(vals.min())
            out["avg_affinity"] = float(vals.mean())
    # Lipinski
    if "lipinski_compliant" in df.columns:
        out["lipinski_rate"] = float((df["lipinski_compliant"].astype(bool).sum() / max(1, len(df))) * 100)
    # Veber/Egan
    if "veber_compliant" in df.columns:
        out["veber_rate"] = float((df["veber_compliant"].astype(bool).sum() / max(1, len(df))) * 100)
    if "egan_compliant" in df.columns:
        out["egan_rate"] = float((df["egan_compliant"].astype(bool).sum() / max(1, len(df))) * 100)
    # QED
    if "qed" in df.columns:
        q = df["qed"].dropna()
        if len(q) > 0:
            out["qed_mean"] = float(q.mean())
            out["qed_median"] = float(q.median())
    # 溶解度等级与毒性
    if "solubility_class" in df.columns:
        out["solubility_counts"] = df["solubility_class"].fillna("未知").value_counts().to_dict()
    if "toxicity_risk_level" in df.columns:
        out["toxicity_counts"] = df["toxicity_risk_level"].fillna("未知").value_counts().to_dict()
    return out


def pick_top10(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    use = df.copy()
    if "binding_affinity" in use.columns:
        return use.sort_values("binding_affinity", ascending=True).head(10)
    return use.head(10)


def add_title_page(document: Document, run_dir: Optional[Path]):
    document.add_heading("PRRSV 抑制剂设计平台 - 项目白皮书", 0)
    p = document.add_paragraph()
    p.add_run(f"版本: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n").bold = True
    if run_dir is not None:
        p.add_run(f"基于最新运行: {run_dir.name}\n")
    p.add_run("Windsurf AI 工程团队 | 机密 - 请勿外传")


def add_section_intro(document: Document):
    document.add_heading("一、项目背景与目标", level=1)
    document.add_paragraph(
        "本项目面向猪繁殖与呼吸综合征病毒（PRRSV）衣壳蛋白与整合素之间的蛋白-蛋白相互作用（PPI）抑制剂发现。"
        "平台集成“深度学习生成 → 分子对接 → ADMET 分析 → 2D/3D 可视化 → 结果管理”的完整工作流，"
        "支持成千上万级分子的批量生成与筛选，并强调可追溯、可复现、可扩展的工程化能力。"
    )
    document.add_paragraph(
        "项目目标在于：1）高效发现潜在先导化合物；2）以可解释的方式对候选化合物进行性质与风险评估；"
        "3）通过工程化的Run管理和自动化报告，支持快速迭代与规模化筛选。"
    )


def add_section_system(document: Document):
    document.add_heading("二、系统功能矩阵", level=1)
    items = [
        "分子生成：支持优化生成、随机生成、CMD-GEN 以及外部 Transformer/Diffusion 挂载。",
        "分子对接：AutoDock Vina 引擎，Meeko 预处理，批量对接并输出 CSV。",
        "ADMET 分析：RDKit 描述符、Lipinski/Veber/Egan 规则、QED、ESOL 近似溶解度、PAINS/SMARTS 毒性警示，缺少 RDKit 时自动回退简化模式。",
        "2D 可视化：单体与网格图，Top-N 上限 1000，并支持 HTML 报告。",
        "3D 可视化：py3Dmol 与 Plotly 交互式仪表盘，蛋白-配体复合物与结合位点分析；无对接结果时回退到配体 CSV。",
        "结果管理：Run 级目录，产物与报告统一归档，可追溯。",
    ]
    for it in items:
        document.add_paragraph(f"• {it}")


def add_section_models(document: Document):
    document.add_heading("三、模型原理与架构", level=1)
    document.add_heading("3.1 深度学习流水线（三阶段）", level=2)
    document.add_paragraph(
        "Phase 1（等变 GNN 评分模型，占位实现）：以蛋白-配体相互作用为背景，构建等变图神经网络用于亲和力或打分的先验估计。"
        "该阶段在代码中提供了占位接口（deep_learning_pipeline.py 的 phase_1_equivariant_gnn），在无训练数据时自动跳过训练并产出占位模型，以保证流水线可运行。"
    )
    document.add_paragraph(
        "Phase 2（Transformer 条件生成）：以 Phase1 的先验为指导，对 SMILES 进行条件生成。平台支持外部 Transformer 推理脚本挂载，通过参数化接口（模块名/函数/权重路径/关键字参数 JSON）实现灵活对接；"
        "内置动态化学空间扩展与规则模板，自动生成 o/m/p 位点二/三取代芳环、苯并唑/五元唑/三嗪/噻二唑/联苯/三联苯/亚当烷/双环等骨架，大幅提升去重后的唯一分子数。"
    )
    document.add_paragraph(
        "Phase 3（Diffusion + RL 优化）：在候选基础上进行扩散模型生成与强化学习（RL）重排/裁剪（占位实现），以多目标（亲和力、QED、规则合规性、溶解度等）为综合目标函数，"
        "实现 Top-K → Final-N 的收敛，产出结构更优、性质更均衡的候选。"
    )

    document.add_heading("3.2 关键创新点", level=2)
    bullets = [
        "化学空间动态扩展：系统化生成与组合多类骨架与取代位点，软上限 ~8000 唯一分子。",
        "强健回退策略：无对接结果则回退配体 CSV；无 RDKit/py3Dmol 则启用简化模式/跳过图像导出，保证全链路不中断。",
        "可插拔深度学习接口：外部 Transformer/Diffusion/RL 模型可随时挂载，统一参数化接口，方便 A/B 测试与快速迁移。",
        "高吞吐可视化：2D Top-N 上限到 1000，3D 仪表盘自动导出 PNG（kaleido），便于报告嵌图与分享。",
        "Run 级工程化管理：标准化目录、步骤记录、结果汇总与复制（深度学习 Phase2/3 CSV 自动复制到 ligands/）。",
    ]
    for b in bullets:
        document.add_paragraph(f"• {b}")

    document.add_heading("3.3 分子对接与准备", level=2)
    document.add_paragraph(
        "采用 AutoDock Vina 进行批量对接，结合 Meeko 进行配体三维构象与加氢；RDKit 提供基础 3D 构象生成与最小化。结果统一输出为 CSV，包括最佳亲和力、RMSD 与成功标记等。"
    )

    document.add_heading("3.4 ADMET 评估方法", level=2)
    document.add_paragraph(
        "在 RDKit 环境下，计算分子量、LogP、HBD/HBA、旋转键、TPSA、环结构、芳香环、摩尔折射率、Csp3 分数、QED 等描述符；"
        "基于 Lipinski/Veber/Egan 规则给出是否合规及违规详情；以 ESOL 近似公式估计溶解度并分级；"
        "以 PAINS 目录与常见 SMARTS 模式标注结构警示并给出风险等级。在无 RDKit 时，使用启发式近似回退仍提供可用的估计。"
    )


def add_section_engineering(document: Document):
    document.add_heading("四、工程设计与可扩展性", level=1)
    document.add_paragraph(
        "Streamlit 统一界面（HJD/unified_web_interface.py）组织全流程；结果管理器（scripts/result_manager.py）以 run_* 目录治理 Run 生命周期，"
        "包含 ligands/、docking/、admet/、visualization_2d/、visualization_3d/、reports/ 等子目录。"
        "2D/3D 可视化组件在缺失对接或依赖时具备回退能力；3D 仪表盘导出 HTML 与 PNG，辅助 README 与白皮书插图。"
    )
    document.add_paragraph(
        "平台设计强调：模块解耦（生成/对接/ADMET/可视化）、数据标准化（CSV/SMILES/列名约定）、可插拔模型接口、以及批处理与可追溯性。"
    )


def add_section_experiments(document: Document, merged_df: pd.DataFrame, images: Dict[str, Optional[Path]]):
    document.add_heading("五、实验设计与结果", level=1)
    document.add_paragraph("本章节基于最近一次 run 的数据，给出核心统计与可视化示例。")

    # 统计
    stats = summarize_numbers(merged_df)
    if stats:
        lines = []
        if "best_affinity" in stats:
            lines.append(f"最佳结合亲和力: {stats['best_affinity']:.2f} kcal/mol")
        if "avg_affinity" in stats:
            lines.append(f"平均结合亲和力: {stats['avg_affinity']:.2f} kcal/mol")
        if "lipinski_rate" in stats:
            lines.append(f"Lipinski 符合率: {stats['lipinski_rate']:.1f}%")
        if "veber_rate" in stats:
            lines.append(f"Veber 符合率: {stats['veber_rate']:.1f}%")
        if "egan_rate" in stats:
            lines.append(f"Egan 符合率: {stats['egan_rate']:.1f}%")
        if "qed_mean" in stats:
            lines.append(f"QED 平均值: {stats['qed_mean']:.3f}")
        if "qed_median" in stats:
            lines.append(f"QED 中位数: {stats['qed_median']:.3f}")
        if lines:
            document.add_paragraph("；".join(lines))

    # 图片区域
    document.add_heading("5.1 可视化图示", level=2)
    safe_add_picture(document, images.get("grid_png"), 6.0, "Top-N 分子 2D 网格图（示例）")
    safe_add_picture(document, images.get("dash_png"), 6.0, "3D 交互式仪表盘静态图（示例）")

    # 备用图（如实际 run 缺图）
    fb = images.get("fallback_imgs", {})
    safe_add_picture(document, fb.get("affinity"), 6.0, "结合亲和力分布（示例）")
    safe_add_picture(document, fb.get("top10"), 6.0, "Top-10 候选分子（示例）")
    safe_add_picture(document, fb.get("lipinski"), 6.0, "Lipinski 合规性（示例）")
    safe_add_picture(document, fb.get("admet_props"), 6.0, "ADMET 性质分布（示例）")

    # Top-10 表格
    if not merged_df.empty:
        document.add_heading("5.2 Top-10 候选（按结合亲和力排序）", level=2)
        top10 = pick_top10(merged_df)
        cols = [c for c in ["compound_id", "smiles", "binding_affinity", "qed", "lipinski_compliant", "solubility_class", "toxicity_risk_level"] if c in top10.columns]
        table = document.add_table(rows=1 + min(10, len(top10)), cols=len(cols))
        hdr_cells = table.rows[0].cells
        for j, c in enumerate(cols):
            hdr_cells[j].text = c
        for i in range(min(10, len(top10))):
            row = table.rows[i + 1].cells
            for j, c in enumerate(cols):
                val = top10.iloc[i][c]
                row[j].text = "" if pd.isna(val) else str(val)


def add_section_conclusion(document: Document):
    document.add_heading("六、结论与展望", level=1)
    document.add_paragraph(
        "平台以工程化与可扩展性为核心，统一深度学习生成、对接与 ADMET 评估流程，"
        "通过强健回退与可插拔接口，在多种环境与数据可用性下均可持续产出结果。"
    )
    document.add_paragraph(
        "后续计划包括：1）对接真实的条件 Transformer 与 Diffusion/RL 模型，开展系统的 A/B 测试；"
        "2）引入更丰富的 ADMET 预测端点（如 pkCSM/DeepTox 等）；3）引入分布式任务编排以进一步提升吞吐；"
        "4）结合实验数据闭环优化。"
    )


def add_section_appendix(document: Document):
    document.add_heading("附录：目录结构与运行说明", level=1)
    items = [
        "results/run_YYYYMMDD_###/ligands/ — 配体 CSV 与 SMILES",
        "results/run_YYYYMMDD_###/docking/ — 对接 CSV 结果",
        "results/run_YYYYMMDD_###/admet/ — ADMET 分析结果",
        "results/run_YYYYMMDD_###/visualization_2d/ — 2D 可视化",
        "results/run_YYYYMMDD_###/visualization_3d/ — 3D 仪表盘与报告",
        "results/run_YYYYMMDD_###/reports/ — 报告与白皮书",
        "HJD/unified_web_interface.py — Streamlit 统一界面",
        "HJD/scripts/admet_analyzer.py — ADMET 分析模块",
        "HJD/scripts/visualization_3d.py — 3D 可视化报告与仪表盘",
        "HJD/deep_learning_pipeline.py — 深度学习流水线（占位 + 外部模型挂载）",
    ]
    for it in items:
        document.add_paragraph(f"• {it}")


def main():
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("--run", type=str, default=None, help="指定 run 目录名，如 run_20250918_008")
    args = parser.parse_args()

    run_dir = find_latest_run_dir(args.run)
    if run_dir is None:
        print("[WARN] 未找到任何 results/run_* 目录，将使用仓库内示例图与通用说明生成白皮书。")

    document = Document()
    add_title_page(document, run_dir)
    add_section_intro(document)
    add_section_system(document)
    add_section_models(document)
    add_section_engineering(document)

    if run_dir is not None:
        data = collect_run_data(run_dir)
        add_section_experiments(document, data.get("merged_df", pd.DataFrame()), data)
        out_dir = data.get("reports_dir", (run_dir / "reports"))
    else:
        # 无 run 的情况下, 使用空DF与示例图生成实验章节
        dummy = {
            "merged_df": pd.DataFrame(),
            "grid_png": EXPERIMENT_REPORT_DIR / "visualization_2d" / "grids" / "Top_50_Molecules_grid.png",
            "dash_png": None,
            "fallback_imgs": {
                "affinity": EXPERIMENT_REPORT_DIR / "binding_affinity_distribution.png",
                "top10": EXPERIMENT_REPORT_DIR / "top_10_molecules.png",
                "lipinski": EXPERIMENT_REPORT_DIR / "lipinski_compliance.png",
                "admet_props": EXPERIMENT_REPORT_DIR / "admet_properties.png",
            },
        }
        add_section_experiments(document, pd.DataFrame(), dummy)
        out_dir = RESULTS_DIR / "run_nodata" / "reports"
        out_dir.mkdir(parents=True, exist_ok=True)

    add_section_conclusion(document)
    add_section_appendix(document)

    outfile = out_dir / "PRRSV_Whitepaper.docx"
    document.save(str(outfile))
    print(f"[OK] 白皮书已生成: {outfile}")

    # 复制一份到 HJD/reports 目录
    try:
        import shutil
        dst = REPORT_OUT_DIR / outfile.name
        shutil.copy2(outfile, dst)
        print(f"[OK] 已复制到: {dst}")
    except Exception as e:
        print(f"[WARN] 拷贝白皮书失败: {e}")


if __name__ == "__main__":
    main()
